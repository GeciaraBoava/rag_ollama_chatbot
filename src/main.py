import sys
import os
from pathlib import Path
from typing import List
import requests
import warnings

# Suprimir avisos
warnings.filterwarnings('ignore', category=DeprecationWarning)
warnings.filterwarnings('ignore', message='.*punkt_tab.*')
warnings.filterwarnings('ignore', message='.*validate_default.*')
warnings.filterwarnings('ignore', category=UserWarning, module='pydantic')
warnings.filterwarnings('ignore', message='.*symlinks.*')
warnings.filterwarnings('ignore', category=UserWarning, module='huggingface_hub')

# Desabilitar warning de symlinks do HuggingFace
os.environ['HF_HUB_DISABLE_SYMLINKS_WARNING'] = '1'

# llama-index 0.11.x (estrutura compatível)
from llama_index.core import Document, VectorStoreIndex, Settings, StorageContext
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.faiss import FaissVectorStore
from llama_index.llms.ollama import Ollama

import faiss
from docx import Document as DocxDocument
from pypdf import PdfReader

DEFAULT_DOCUMENTS_FOLDER = "documentos"
EMBEDDING_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
OLLAMA_MODEL_NAME = "llama3.2:3b"  # Modelo mais rápido
EXIT_COMMANDS = ["sair", "exit", "quit"]

# Configurações otimizadas para velocidade + precisão
CHUNK_SIZE = 512  # Captura contexto suficiente
CHUNK_OVERLAP = 50  # Boa continuidade
SIMILARITY_TOP_K = 4  # Reduzido de 5 para 4 (menos contexto = mais rápido)
MAX_TOKENS = 800  # Reduzido de 1024 (respostas mais diretas)
NUM_CTX = 3072  # Reduzido de 4096 (menos contexto = mais rápido)
FAISS_INDEX_DIR = "./storage"  # Diretório para persistir índice
SYSTEM_PROMPT = """Responda em português de forma direta e objetiva.
Liste todos os itens mencionados no contexto.
Seja breve e preciso."""

def check_ollama_running():
    """Verifica se o Ollama está rodando."""
    try:
        resp = requests.get("http://localhost:11434/api/tags", timeout=5)
        return resp.status_code == 200
    except Exception:
        return False


def check_gpu_available():
    """Verifica se GPU está disponível."""
    try:
        import torch
        if torch.cuda.is_available():
            gpu_name = torch.cuda.get_device_name(0)
            print(f"🎮 GPU detectada: {gpu_name}")
            return True
        else:
            print("💻 Usando CPU (sem GPU detectada)")
            return False
    except Exception:
        print("💻 Usando CPU")
        return False

def load_documents(folder_path: str) -> List[Document]:
    """Carrega documentos da pasta especificada."""
    folder = Path(folder_path)
    if not folder.exists() or not folder.is_dir():
        raise FileNotFoundError(f"A pasta '{folder_path}' não existe.")
    
    files = list(folder.glob("*"))
    if not files:
        raise ValueError(f"A pasta '{folder_path}' está vazia.")

    docs = []
    for f in files:
        try:
            if f.suffix.lower() == ".txt":
                content = f.read_text(encoding="utf-8")
            elif f.suffix.lower() == ".docx":
                docx = DocxDocument(str(f))
                content = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])
            elif f.suffix.lower() == ".pdf":
                reader = PdfReader(str(f))
                content = "\n".join([p.extract_text() or "" for p in reader.pages])
            else:
                print(f"⚠️  Ignorando {f.name} (tipo não suportado)")
                continue

            if content and content.strip():
                docs.append(
                    Document(
                        text=content, 
                        doc_id=str(f.name), 
                        metadata={"source": str(f)}
                    )
                )
                print(f"✅ {f.name} carregado ({len(content)} chars)")
            else:
                print(f"⚠️  Arquivo vazio: {f.name}")
        except Exception as e:
            print(f"❌ Erro lendo {f.name}: {e}")

    if not docs:
        raise ValueError("Nenhum documento válido encontrado.")
    
    print(f"✅ Documentos carregados: {len(docs)} arquivo(s).")
    return docs

def create_faiss_vector_store(embedding_dim: int = 384):
    """Cria um vector store FAISS."""
    index = faiss.IndexFlatL2(embedding_dim)
    return FaissVectorStore(faiss_index=index)

def setup_rag_system(documents_folder: str = DEFAULT_DOCUMENTS_FOLDER):
    """Configura o sistema RAG completo."""
    print("📦 Configurando RAG...")
    
    # Verifica se existe índice salvo
    if Path(FAISS_INDEX_DIR).exists():
        print("🔄 Carregando índice existente...")
        try:
            from llama_index.core import load_index_from_storage
            
            # Configurar embeddings e LLM antes de carregar
            embed_model = HuggingFaceEmbedding(
                model_name=EMBEDDING_MODEL_NAME,
                cache_folder="./.cache/embeddings"
            )
            
            llm = Ollama(
                model=OLLAMA_MODEL_NAME, 
                request_timeout=60.0,
                temperature=0.0,  # Zero para respostas mais diretas
                additional_kwargs={
                    "num_predict": MAX_TOKENS,
                    "num_ctx": NUM_CTX,
                    "num_thread": 8,  # Usar múltiplas threads
                },
                system_prompt = SYSTEM_PROMPT
            )
            
            Settings.llm = llm
            Settings.embed_model = embed_model
            Settings.chunk_size = CHUNK_SIZE
            Settings.chunk_overlap = CHUNK_OVERLAP
            
            storage_context = StorageContext.from_defaults(persist_dir=FAISS_INDEX_DIR)
            index = load_index_from_storage(storage_context)
            
            query_engine = index.as_query_engine(
                similarity_top_k=SIMILARITY_TOP_K,
                streaming=False,
                response_mode="compact"
            )
            print("✅ Índice carregado com sucesso!")
            return query_engine
        except Exception as e:
            print(f"⚠️  Erro ao carregar índice: {e}")
            print("🔄 Criando novo índice...")
    
    # Criar novo índice
    documents = load_documents(documents_folder)

    print("🔧 Configurando embeddings...")
    embed_model = HuggingFaceEmbedding(
        model_name=EMBEDDING_MODEL_NAME,
        cache_folder="./.cache/embeddings"
    )

    print("🔧 Configurando LLM (Ollama)...")
    llm = Ollama(
        model=OLLAMA_MODEL_NAME, 
        request_timeout=60.0,
        temperature=0.0,  # Zero para respostas mais diretas
        additional_kwargs={
            "num_predict": MAX_TOKENS,
            "num_ctx": NUM_CTX,
            "num_thread": 8,  # Usar múltiplas threads
        },
        system_prompt = SYSTEM_PROMPT
    )

    Settings.llm = llm
    Settings.embed_model = embed_model
    Settings.chunk_size = CHUNK_SIZE
    Settings.chunk_overlap = CHUNK_OVERLAP

    print("🔧 Configurando FAISS...")
    vector_store = create_faiss_vector_store(embedding_dim=384)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    print("🔧 Criando índice...")
    index = VectorStoreIndex.from_documents(
        documents,
        storage_context=storage_context,
        show_progress=True
    )

    # Salvar índice para uso futuro
    print("💾 Salvando índice...")
    index.storage_context.persist(persist_dir=FAISS_INDEX_DIR)

    query_engine = index.as_query_engine(
        similarity_top_k=SIMILARITY_TOP_K,
        streaming=False,
        response_mode="compact"  # Compact é mais rápido que tree_summarize
    )
    print("✅ Sistema RAG configurado e salvo!")
    return query_engine

def run_chat_loop(query_engine):
    """Loop principal do chatbot."""
    print("\n" + "="*40)
    print("🤖 Chatbot RAG — digite sua pergunta")
    print("Digite 'sair' para encerrar.")
    print("="*40 + "\n")
    
    import time
    
    while True:
        try:
            pergunta = input("👤 Você: ").strip()
            if not pergunta:
                continue
            if pergunta.lower() in EXIT_COMMANDS:
                print("👋 Encerrando.")
                break
            
            print("🔍 Processando...")
            start_time = time.time()
            
            resposta = query_engine.query(pergunta)
            
            elapsed = time.time() - start_time
            print(f"\n🤖 Bot: {resposta}")
            print(f"⏱️  Tempo: {elapsed:.2f}s\n")
        except KeyboardInterrupt:
            print("\n👋 Interrompido pelo usuário.")
            break
        except Exception as e:
            print(f"❌ Erro: {e}")

def main():
    """Função principal."""
    print("="*60)
    print("🤖 RAG Chatbot - Inicializando...")
    print("="*60 + "\n")
    
    # Verifica se o modelo está disponível
    print(f"📋 Verificando modelo: {OLLAMA_MODEL_NAME}")
    if not check_ollama_running():
        print("❌ Ollama não está rodando. Rode: `ollama serve`")
        sys.exit(1)
    
    # Verifica se o modelo está instalado
    try:
        import ollama
        models = ollama.list()
        model_names = [m['name'] for m in models.get('models', [])]
        
        if not any(OLLAMA_MODEL_NAME in name for name in model_names):
            print(f"⚠️  Modelo '{OLLAMA_MODEL_NAME}' não encontrado.")
            print(f"📥 Baixando modelo... (isso pode levar alguns minutos)")
            ollama.pull(OLLAMA_MODEL_NAME)
            print(f"✅ Modelo '{OLLAMA_MODEL_NAME}' baixado!")
        else:
            print(f"✅ Modelo '{OLLAMA_MODEL_NAME}' disponível!")
    except Exception as e:
        print(f"⚠️  Não foi possível verificar modelo: {e}")
        print("   Continuando mesmo assim...")

    try:
        query_engine = setup_rag_system()
        run_chat_loop(query_engine)
    except Exception as e:
        print(f"❌ Erro crítico: {e}")
        raise

if __name__ == "__main__":
    main()